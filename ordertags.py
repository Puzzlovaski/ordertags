from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import re

def add_logo(p):
    p.add_run().add_picture('./logo.jpg',width = Inches(2))
    p.add_run(text = '\n')

document = Document('./template.docx')
style = document.styles['Normal']
font = style.font
font.size = Pt(12)
font.name = 'Leelawadee UI'

photos = open("photo.txt", "w", encoding="utf8")
with open("input.txt", "r", encoding="utf8") as f:
    lines = f.readlines()
    p = document.add_paragraph('')
    p.style = document.styles['Normal']
    p.paragraph_format.keep_together = True
    order = 0
    enter = False
    first_enter = False
    for l in lines:
        print(l)
        if l == '\n':
            if first_enter:
                first_enter = False
                run = p.add_run(text = l)
        elif re.search("^[0-9][0-9]:[0-9][0-9].*\[Photo\]$",l) or re.search("^[0-9]:[0-9][0-9].*\[Photo\]$",l):
            photos.write(l)
        elif re.search("^[0-9][0-9]:[0-9][0-9]",l) or re.search("^[0-9]:[0-9][0-9]",l):
            order += 1
            p = document.add_paragraph('__________________________________'+str(order)+'\n')
            p.style = document.styles['Normal']
            p.paragraph_format.keep_together = True
            run = p.add_run(text = l)
            enter = False
            first_enter = True
        elif "ปลายทาง" in l :
            run = p.add_run(text = l)
            f = p.add_run(text= '\nCOD')
            f.font.highlight_color = WD_COLOR.RED
            f.bold = True
        elif "ผู้รับ" in l:
            add_logo(p)
            run = p.add_run(text= l).bold = True
        else :
            run = p.add_run(text= l)

document.save('output.docx')
